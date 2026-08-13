"""
Word documents.

Three, for three different readers:

* **Company profile** — what to send someone who asks "what do you do?"
* **Business plan** — the internal planning document, with the figures we do
  not hold left blank rather than guessed.
* **Onboarding handbook** — the non-engineering half of a new joiner's first
  week. The engineering half is `Docs/00-start-here.md`, and this document
  points at it rather than restating it.

Everything factual comes from the export.
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .brand import (
    CYAN, VIOLET, INK, SLATE, MUTED, LINE, PAPER_ALT, FONT_SANS, FONT_SERIF,
    TBD, PLACEHOLDER_NOTE, ordered_categories, rupees, footer_line,
    generated_stamp,
)


def _rgb(h: str) -> RGBColor:
    return RGBColor.from_string(h)


def _shade(cell, hexcolor: str):
    """Cell background. python-docx has no API for this; it needs raw XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _base_document(data: dict, title: str, subtitle: str) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT_SANS
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(SLATE)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Running footer, so a printed page that gets separated from the stack can
    # still be traced back to a company and a generation date.
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(footer_line(data) + "\n" + generated_stamp(data))
    run.font.size = Pt(7.5)
    run.font.color.rgb = _rgb(MUTED)
    run.font.name = FONT_SANS

    _cover(doc, data, title, subtitle)
    return doc


def _cover(doc: Document, data: dict, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(data["company"]["name"].upper())
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = _rgb(CYAN)
    r.font.name = FONT_SANS

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = _rgb(INK)
    r.font.name = FONT_SERIF

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = _rgb(MUTED)
    r.font.name = FONT_SANS

    _rule(doc)


def _rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), CYAN)
    bdr.append(bottom)
    pPr.append(bdr)


def _h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = _rgb(INK)
    r.font.name = FONT_SERIF
    return p


def _h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = _rgb(CYAN)
    r.font.name = FONT_SANS
    return p


def _para(doc: Document, text: str, *, size=10.5, color=SLATE, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    r.font.italic = italic
    r.font.name = FONT_SANS
    return p


def _bullet(doc: Document, text: str):
    """Bold lead-in before an em dash, so a long list stays scannable."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if " \u2014 " in text:
        head, tail = text.split(" \u2014 ", 1)
        r1 = p.add_run(head + " \u2014 ")
        r1.font.bold = True
        r1.font.color.rgb = _rgb(INK)
        r1.font.size = Pt(10.5)
        r1.font.name = FONT_SANS
        r2 = p.add_run(tail)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = _rgb(SLATE)
        r2.font.name = FONT_SANS
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = _rgb(SLATE)
        r.font.name = FONT_SANS
    return p


def _table(doc: Document, headers, rows, *, widths=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], INK)
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = _rgb("FFFFFF")
        r.font.name = FONT_SANS

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[ci], PAPER_ALT)
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(size)
            r.font.color.rgb = _rgb(SLATE)
            r.font.name = FONT_SANS

    if widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _callout(doc: Document, text: str):
    """A single-cell shaded table. Word has no boxed-note primitive."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    _shade(cell, PAPER_ALT)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = _rgb(SLATE)
    r.font.name = FONT_SANS
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ------------------------------------------------------------ company profile

def build_company_profile(data: dict, out_path):
    c, cat, com, eng = (data["company"], data["catalogue"],
                        data["commercial"], data["engineering"])

    doc = _base_document(
        data, "Company Profile",
        "Connected hardware, designed and built end to end.",
    )

    _h1(doc, "Who we are")
    _para(doc,
          f"{c['name']} designs, manufactures and operates connected hardware for homes "
          f"and small businesses. We are based in {c['location']} and sell directly at "
          f"{c['site']}.")
    _para(doc,
          "We are unusual in this category in that we own every layer: the circuit boards, "
          "the firmware inside them, the platform they connect to, the apps that control "
          "them, and the shop that sells them. Most competitors resell a white-label device "
          "with their own app on top. Owning the whole stack is why a fault found in the "
          "morning can reach a customer's device the same day.")

    _h1(doc, "What we build")
    _bullet(doc, f"Hardware \u2014 {eng['hardwareProjects']} board-level projects, our own designs.")
    _bullet(doc, f"Firmware \u2014 {eng['firmwareDeviceTypes']} device types sharing one library, "
                 "so provisioning, connectivity and over-the-air update behave identically "
                 "across the fleet.")
    _bullet(doc, "Platform \u2014 a self-hosted control plane that owns devices, telemetry, "
                 "commands, scenes and automations.")
    _bullet(doc, "Applications \u2014 a web console and a native mobile app, kept at feature parity.")
    _bullet(doc, "Commerce \u2014 our own storefront, checkout, order handling and support.")

    _h1(doc, "Product range")
    _para(doc,
          f"{cat['total']} products across {len(cat['categories'])} categories, "
          f"from {rupees(cat['priceMin'])} to {rupees(cat['priceMax'])}.")
    rows = []
    for name in ordered_categories(cat):
        items = [p for p in cat["products"] if p["category"] == name]
        rows.append([
            name, str(len(items)),
            f"{rupees(min(p['price'] for p in items))} \u2013 "
            f"{rupees(max(p['price'] for p in items))}",
            ", ".join(p["name"].replace("Circuvent ", "") for p in items[:4])
            + (" …" if len(items) > 4 else ""),
        ])
    _table(doc, ["Category", "Products", "Price range", "Examples"], rows,
           widths=[1.4, 0.8, 1.5, 2.8])

    _h1(doc, "How our products work together")
    _para(doc,
          "Every device speaks the same protocol to the same platform, so a customer who "
          "starts with a single smart plug can add a water controller, a camera and a lock "
          "without a second app or a second account. Scenes and automations work across "
          "categories: a motion sensor can drive a light, a gate can trigger a camera "
          "recording, a tank level can start a pump.")
    _para(doc,
          "Commands reach hardware in well under a second, and nothing in that path polls. "
          "Devices hold sensible local behaviour when the network is unavailable, and "
          "recover on their own when it returns.")

    _h1(doc, "Commercial terms")
    _table(doc, ["Term", "Detail"], [
        ["Warranty", f"{com['warrantyMonths']} months on every device"],
        ["Returns", f"{com['returnDays']} days from delivery, unused and in original packaging"],
        ["Free shipping", f"Orders over {rupees(com['freeShippingOver'])}"],
        ["Standard shipping", f"{rupees(com['flatShipping'])} flat"],
        ["Bundles", "Multi-device bundles discounted automatically at checkout"],
        ["Support", c["supportEmail"]],
    ], widths=[1.8, 4.7])
    _callout(doc,
             "These terms are generated from the same source the website and every customer "
             "invoice read. They cannot drift from what a customer is quoted at checkout.")

    _h1(doc, "Contact")
    _table(doc, ["", ""], [
        ["Website", c["site"]],
        ["Sales", c["salesEmail"]],
        ["Support", c["supportEmail"]],
        ["Phone", c["phone"]],
        ["Location", c["location"]],
    ], widths=[1.5, 5.0])

    doc.save(str(out_path))
    return out_path


# --------------------------------------------------------------- business plan

def build_business_plan(data: dict, out_path):
    c, cat, com, eng = (data["company"], data["catalogue"],
                        data["commercial"], data["engineering"])

    doc = _base_document(
        data, "Business Plan",
        "Internal planning document — not for external distribution.",
    )

    _callout(doc, PLACEHOLDER_NOTE)

    _h1(doc, "1. Executive summary")
    _para(doc,
          f"{c['name']} builds connected hardware for Indian homes and small businesses, "
          f"selling {cat['total']} products directly through {c['site']} at prices from "
          f"{rupees(cat['priceMin'])} to {rupees(cat['priceMax'])}. Unlike resellers of "
          "white-label hardware, we design the boards, write the firmware, run the platform "
          "and operate the storefront ourselves.")
    _para(doc,
          "The platform is built and in production: "
          f"{eng['firmwareDeviceTypes']} firmware device types, "
          f"{eng['hardwareProjects']} hardware projects and {eng['deployables']} independently "
          "deployed systems. The immediate opportunity is distribution — the engineering "
          "is further ahead than the channel.")

    _h1(doc, "2. The problem")
    _bullet(doc, "Unreliable connectivity \u2014 devices that drop off the network and need a "
                 "physical reset to return, often because the device boots faster than the router.")
    _bullet(doc, "Silent failure \u2014 apps offering controls the hardware does not implement, "
                 "so a customer reports faulty hardware for a software defect.")
    _bullet(doc, "Fragmented ownership \u2014 hardware from one vendor and an app from another, "
                 "with neither accountable when it fails.")
    _bullet(doc, "No record \u2014 when a pump burns out or a gate opens at night, "
                 "there is no history to consult.")

    _h1(doc, "3. Our solution")
    _para(doc,
          "One platform covering every product we sell, with the whole stack under our "
          "control. Concretely, that gives us four things competitors assembling parts "
          "cannot easily match:")
    _bullet(doc, "Same-day fixes \u2014 a firmware defect can be diagnosed, fixed and "
                 "delivered over the air without a vendor in the loop.")
    _bullet(doc, "Cross-category automation \u2014 a sensor in one category can drive a "
                 "device in another, because they share one engine.")
    _bullet(doc, "Unattended recovery \u2014 devices retry indefinitely and only ask for "
                 "setup when genuinely unconfigured, so a power cut does not generate support calls.")
    _bullet(doc, "Consistent terms \u2014 warranty, returns and delivery are quoted from a "
                 "single source, so the invoice and the policy page cannot disagree.")

    _h1(doc, "4. Products and pricing")
    # Same category order as the customer-facing documents, cheapest first
    # within each, so the entry price leads.
    plan_rank = {name: i for i, name in enumerate(ordered_categories(cat))}
    rows = [[p["name"], p["category"], rupees(p["price"]),
             rupees(p["compareAt"]) if p["compareAt"] else "\u2014",
             f"{p['discountPct']}%" if p["discountPct"] else "\u2014"]
            for p in sorted(cat["products"],
                            key=lambda x: (plan_rank.get(x["category"], 99), x["price"]))]
    _table(doc, ["Product", "Category", "Price", "MRP", "Saving"], rows,
           widths=[2.2, 1.5, 0.9, 0.9, 0.7], size=8.5)

    _h1(doc, "5. Market")
    _para(doc,
          "Indian home automation is growing from a low base, driven by cheaper connectivity, "
          "rising electricity costs and wider smartphone use. Our positioning is deliberately "
          "practical rather than aspirational: water management, energy monitoring, safety "
          "and access are recurring household problems with a measurable cost, and they are "
          "under-served relative to lighting and entertainment.")
    _table(doc, ["Market input", "Value"], [
        ["Total addressable market", TBD],
        ["Serviceable market", TBD],
        ["Current market share", TBD],
        ["Primary competitors", TBD],
        ["Average selling price", rupees(sum(p["price"] for p in cat["products"]) / cat["total"])],
    ], widths=[3.0, 3.5])

    _h1(doc, "6. Go to market")
    _h2(doc, "Today")
    _bullet(doc, "Direct online sales \u2014 our own storefront, checkout and fulfilment.")
    _bullet(doc, "Bundles \u2014 multi-device kits that raise order value and platform stickiness.")
    _h2(doc, "Next")
    _bullet(doc, "Installer and electrician channel \u2014 the people already in the customer's "
                 "home when the decision is made.")
    _bullet(doc, "Builder and developer tie-ups \u2014 fitted at construction rather than retrofit.")
    _bullet(doc, "Retail distribution \u2014 physical presence for a category customers "
                 "still like to see.")
    _bullet(doc, "Commercial deployments \u2014 multi-site access control and monitoring, "
                 "where per-site value is higher.")

    _h1(doc, "7. Operations")
    _table(doc, ["Function", "Position today"], [
        ["Hardware design", f"In-house — {eng['hardwareProjects']} board projects"],
        ["Firmware", f"In-house — {eng['firmwareDeviceTypes']} device types, one shared library"],
        ["Manufacturing", TBD],
        ["Assembly and flashing", "Devices ship with our firmware pre-flashed"],
        ["Fulfilment", TBD],
        ["Support", f"Direct, {c['supportEmail']}"],
        ["Platform hosting", "Self-hosted control plane; storefront on managed hosting"],
    ], widths=[1.8, 4.7])

    _h1(doc, "8. Financial plan")
    _para(doc,
          "Trading figures are not held in the product repository this document is generated "
          "from. They are listed here so the plan is structurally complete, and must be "
          "filled in by the finance owner before circulation.")
    _table(doc, ["Line", "Current", "Year 1", "Year 2", "Year 3"], [
        ["Revenue", TBD, TBD, TBD, TBD],
        ["Units sold", TBD, TBD, TBD, TBD],
        ["Gross margin %", TBD, TBD, TBD, TBD],
        ["Operating cost", TBD, TBD, TBD, TBD],
        ["Headcount", TBD, TBD, TBD, TBD],
        ["Net position", TBD, TBD, TBD, TBD],
    ], widths=[1.5, 1.25, 1.25, 1.25, 1.25], size=8.5)

    _h1(doc, "9. Risks")
    _table(doc, ["Risk", "Mitigation"], [
        ["Component supply and pricing",
         "Multiple footprints where possible; one shared firmware library reduces the cost "
         "of substituting a part"],
        ["Platform outage",
         "Systems deploy independently; devices retain local behaviour and the shop stays up "
         "if the device platform is down"],
        ["A firmware defect that strands devices",
         "The worst case in this business: an unreachable device cannot receive its own fix. "
         "Connectivity changes get the most conservative review of anything we ship"],
        ["Distribution dependence on direct online sales",
         "Installer, retail and builder channels are the main commercial priority"],
        ["Key-person concentration", TBD],
        ["Regulatory and certification", TBD],
    ], widths=[2.3, 4.2])

    _h1(doc, "10. Funding")
    _table(doc, ["Item", "Value"], [
        ["Raised to date", TBD],
        ["Current runway", TBD],
        ["Amount sought", TBD],
        ["Use of funds", TBD],
        ["Valuation expectation", TBD],
    ], widths=[2.5, 4.0])

    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------- onboarding

def build_onboarding_handbook(data: dict, out_path):
    c, cat, com, eng = (data["company"], data["catalogue"],
                        data["commercial"], data["engineering"])

    doc = _base_document(
        data, "New Joiner Handbook",
        "Everything to know in your first week that is not code.",
    )

    _h1(doc, "Welcome")
    _para(doc,
          f"You have joined a company that builds connected hardware end to end \u2014 "
          f"{cat['total']} products, {eng['hardwareProjects']} board designs and "
          f"{eng['firmwareDeviceTypes']} firmware device types, all of it ours. "
          "This handbook covers the company, the product and how we work. "
          "It does not cover setting up your development environment.")
    _callout(doc,
             "Engineers: your environment setup is Docs/00-start-here.md in the repository. "
             "It gets all four systems running on your machine in about half a day, and it "
             "is deliberately not duplicated here \u2014 a second copy of setup instructions "
             "is a second copy to go stale.")

    _h1(doc, "What the company does")
    _para(doc,
          f"{c['name']} sells connected devices for homes and small businesses directly at "
          f"{c['site']}, from {c['location']}. We design the circuit boards, write the "
          "firmware, run the platform, build both apps and operate the shop.")
    _para(doc,
          "The practical consequence, and the thing worth internalising in week one: there "
          "is nobody else to escalate to. If a customer's device misbehaves, the cause is "
          "in our hardware, our firmware, our platform or our app. That is a responsibility "
          "and also the reason we can fix things quickly.")

    _h1(doc, "The product, in one page")
    _table(doc, ["Category", "Products", "What it solves"], [
        ["Home Automation",
         str(cat["categoryCounts"].get("Home Automation", 0)),
         "Switching, dimming, fans, curtains and scenes"],
        ["Safety",
         str(cat["categoryCounts"].get("Safety", 0)),
         "Cameras, locks, gates, motion and number-plate recognition"],
        ["Water Management",
         str(cat["categoryCounts"].get("Water Management", 0)),
         "Tank levels, pump automation and dry-run protection"],
        ["Energy",
         str(cat["categoryCounts"].get("Energy", 0)),
         "Live consumption monitoring"],
    ], widths=[1.6, 0.9, 4.0])

    _h1(doc, "The four systems")
    _para(doc, "Whatever your role, these names come up daily:")
    _bullet(doc, "Devices \u2014 the hardware, running our firmware.")
    _bullet(doc, "Control plane \u2014 the server that owns devices, telemetry and automations.")
    _bullet(doc, "Applications \u2014 the web console and the mobile app.")
    _bullet(doc, "Shop \u2014 the storefront, orders and support.")
    _para(doc,
          "They deploy independently and deliberately do not share a database. "
          "If the shop is down, devices keep working. If the device platform is down, "
          "the shop keeps selling.")

    _h1(doc, "How we work")
    _h2(doc, "Make failure loud")
    _para(doc,
          "The bugs that survive here are the quiet ones \u2014 a control that does nothing, "
          "a setting that saves and vanishes, an action reported as sent that never was. "
          "Nothing errors and nothing logs, so the customer reports broken hardware. "
          "Whatever your role, prefer the version that fails visibly and early over the "
          "version that degrades in silence.")
    _h2(doc, "One owner per fact")
    _para(doc,
          "Anything quoted twice will disagree eventually. Prices, warranty terms, the "
          "support address \u2014 each has exactly one home, and everything else reads it. "
          "This handbook is generated from that same source, which is why the figures in "
          "it match the website.")
    _h2(doc, "Write down why, not what")
    _para(doc,
          "In code, in tickets and in documents: the decision and the constraint are what "
          "someone needs six months later. What was done is already visible.")

    _h1(doc, "Customer terms you should know")
    _table(doc, ["Term", "Detail"], [
        ["Warranty", f"{com['warrantyMonths']} months"],
        ["Returns", f"{com['returnDays']} days, unused and in original packaging"],
        ["Free shipping", f"Over {rupees(com['freeShippingOver'])}"],
        ["Standard shipping", rupees(com["flatShipping"])],
        ["Support address", c["supportEmail"]],
    ], widths=[1.8, 4.7])
    _callout(doc,
             "Never quote a customer terms from memory. These are the published ones and "
             "they are what the invoice says.")

    _h1(doc, "Your first week")
    _bullet(doc, "Day 1 \u2014 accounts, hardware, introductions. Engineers: work through "
                 "Docs/00-start-here.md and get all four systems running.")
    _bullet(doc, "Day 2 \u2014 use the product. Set up real devices, break them, watch what "
                 "the app does. Nothing substitutes for this.")
    _bullet(doc, "Day 3 \u2014 read Docs/01-architecture.md and Docs/26-glossary.md. "
                 "Learn how a command reaches a device.")
    _bullet(doc, "Day 4 \u2014 pick up a first task. Engineers: Docs/27-first-tasks.md is a "
                 "graded ladder, start at level one even if it looks easy.")
    _bullet(doc, "Day 5 \u2014 ship something small and review it with someone.")

    _h1(doc, "Asking for help")
    _para(doc,
          "Ask early. Blocked for an hour on setup is normal; blocked for a day without "
          "asking is not. When you ask, include what you ran, what you expected and what "
          "happened \u2014 that turns a conversation into a fix.")
    _para(doc,
          "If you find a document that is wrong, fix it. These documents are written from "
          "the code and stay true only because the person who notices the drift corrects it. "
          "That includes on your first week \u2014 especially on your first week, when you "
          "are the only person who can still see what is confusing.")

    _h1(doc, "Contacts")
    _table(doc, ["", ""], [
        ["Website", c["site"]],
        ["Sales", c["salesEmail"]],
        ["Support", c["supportEmail"]],
        ["Phone", c["phone"]],
        ["Office", c["location"]],
        ["HR / people", TBD],
        ["IT and accounts", TBD],
    ], widths=[1.5, 5.0])

    doc.save(str(out_path))
    return out_path
