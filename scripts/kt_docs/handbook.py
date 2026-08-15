"""
The knowledge-transfer handbook.

The deck is for a session with somebody talking over it. This is the same
material for somebody reading alone at 9pm, which is when most handover material
is actually read.

It reuses the body helpers from the business documents — headings, tables,
callouts — so a handbook and a company profile printed on the same day look like
they came from the same company. What it does not reuse is the footer: those
documents are generated "from the live product catalogue", and this one is
generated from the repository, which is a different claim and should not be
borrowed.

WHAT IT DELIBERATELY DOES NOT DO

It does not restate `Docs/`. Twenty-nine documents already explain the system in
more detail than a handbook should, they live beside the code, and they are
updated by whoever notices they are wrong. A handbook that copied them would be
a second, worse, un-maintained set. So this points, repeatedly and by filename.
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from business_docs.brand import FONT_SANS, MUTED, SLATE, footer_line
from business_docs.documents import (
    Document, _rgb, _cover, _rule, _h1, _h2, _para, _bullet, _table, _callout,
)
from docx.shared import Inches


def _kt_document(data: dict, title: str, subtitle: str) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT_SANS
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(SLATE)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(
        footer_line(data) + "\n"
        + f"Generated {data['generatedDate']} from the repository at {data['commit']} "
        + "· Docs/ is the source of truth"
    )
    run.font.size = Pt(7.5)
    run.font.color.rgb = _rgb(MUTED)
    run.font.name = FONT_SANS

    _cover(doc, data, title, subtitle)
    return doc


def build_kt_handbook(data: dict, out_path) -> int:
    c = data["counts"]
    doc = _kt_document(
        data,
        "Engineering Handover",
        "What to know before you change anything",
    )

    # ------------------------------------------------------------ orientation
    _h1(doc, "How to use this")
    _para(doc,
          "This handbook is an index with opinions. The system is documented in "
          f"{c['docs']} files under Docs/, written from the code and kept beside it, and "
          "those are the source of truth. What a new engineer lacks is not detail — it is "
          "knowing which of them matters today, and the handful of facts that are true "
          "across all four deployables and therefore live in none of them.")
    _para(doc,
          "Nothing in this document is typed by hand. The device list is the firmware tree, "
          "the document index is Docs/, the traps table is parsed out of "
          "Docs/00-start-here.md, and the counts are counted. If a number here is wrong, "
          "the repository changed and this document has not been rebuilt — run "
          "npm run docs:kt.")
    _callout(doc,
             "Read Docs/00-start-here.md first and all the way through. It gets all four "
             "deployables running on your machine and proves each one works. Budget half a "
             "day. Everything below assumes you have done it.")

    # ------------------------------------------------------------ the system
    _h1(doc, "The system, in one page")
    _para(doc,
          "Circuvent is four deployables that talk over documented contracts. They can be "
          "worked on, deployed and broken independently, and most tasks touch exactly one.")
    _table(
        doc,
        ["Deployable", "Lives in", "Runs on", "Owns"],
        [[d.name, d.path, d.runs_on, d.owns] for d in data["deployables"]],
        widths=[1.4, 0.8, 1.5, 2.8],
    )

    _h2(doc, "There are two databases")
    _para(doc,
          "This surprises everybody, so it is worth stating plainly: the shop and the "
          "control plane do not share a database. The shop stores customers, orders and "
          "wallets in a Neon Postgres reached from the Next.js app, hashing passwords with "
          "scrypt. The control plane stores fleet users, devices and telemetry in a Postgres "
          "container on the VM, hashing with bcrypt.")
    _para(doc,
          "Neither can see the other, and that is deliberate: a compromise or an outage of "
          "the shop cannot reach the device fleet, and vice versa. They are joined only by a "
          "single-sign-on bridge that lets each side vouch for a customer it has already "
          "authenticated. Almost every confused question about users traces back to assuming "
          "there is one. See Docs/05-databases.md.")

    _h2(doc, "How a command reaches a device")
    for step in [
        "The console or app calls POST /devices/:id/command on the control plane.",
        "The API checks the caller's JWT and that the caller owns that device.",
        "The API publishes the body to cv/<deviceId>/cmd on Mosquitto at QoS 1.",
        "The device applies the change and republishes its full state, retained.",
        "The API persists that state and pushes it to every WebSocket client that owns the device.",
        "The UI, which had already applied the change optimistically, reconciles.",
    ]:
        _bullet(doc, step)
    _para(doc,
          "Typical end-to-end latency is well under a second, and there is no polling "
          "anywhere in that path. The device is the authority on its own state — the "
          "interface never assumes it won. Camera frames are the exception: they ride a "
          "dedicated topic, are never persisted and are never retained, because a 10 fps "
          "camera would otherwise write 36,000 rows an hour and a retained frame would hand "
          "the last picture taken to anything that subscribed later.")

    # ------------------------------------------------------------- the fleet
    _h1(doc, "The fleet")
    _para(doc,
          f"The firmware tree ships {c['devices']} device types. Every one of them must have "
          "a control surface in the console — tests/firmware-console-parity.test.ts fails if "
          "one does not, because the failure it prevents is silent: a new sketch lands, the "
          "hardware works, and the console falls through to a generic renderer with a raw "
          "state dump and no controls.")
    devices = data["devices"]
    _para(doc, ", ".join(devices) + ".", size=10)
    _callout(doc,
             "A device type id is the worst thing in this codebase to get wrong. The same "
             "string appears in firmware, the API, both apps and the shop. Copy it, never "
             "retype it — and read Docs/07-adding-a-new-device.md before adding one, because "
             "a type is registered in several tables and adding it to one produces a device "
             "with no controls and no error.")

    # ------------------------------------------------------- how work is done
    _h1(doc, "How work is done here")
    _h2(doc, "The rule that matters most: parity")
    _para(doc,
          "Almost every real bug in this codebase has been one surface knowing something "
          "another surface did not. Not a crash, not a failed request — a control that is "
          "simply absent, or a button that does nothing, on one screen out of six.")
    for ex in [
        "A device type was registered in the phone app but not in the console's type table. "
        "Customers could dim a lamp on their phone and not in a browser. Nothing errored.",
        "The scene editor offered a Power toggle for a curtain. The command builder refuses "
        "power on a curtain and the editor drops actions with no command, so the row was "
        "offered, configured, saved, and silently vanished.",
        "A camera streamed at three frames a second because one resolution setting governed "
        "both stills and video, and nothing on screen connected the two.",
    ]:
        _bullet(doc, ex)
    _para(doc,
          "They share a shape: two tables that must agree, no mechanism forcing them to, and "
          "no error when they disagree. So when you add a fact about a device — a type, a "
          "field, a control, an icon, a label — search for every table that already stores "
          "that kind of fact before you add it to one, then either make one derive from the "
          "other or add a test that fails when they disagree. tests/ is full of that pattern; "
          "device-type-parity.test.ts is the model.")

    _h2(doc, "Make the failure loud")
    _para(doc,
          "The bugs that survive here are the quiet ones. Prefer a design that breaks at "
          "build or test time over one that degrades at runtime. A missing device type should "
          "fail a test, not render a generic chip. An action that builds no command should be "
          "refused, not dropped. An operation that could not be performed must not be counted "
          "as performed. And a refusal and a fault need different words, because a fault "
          "invites a retry and a refusal will never succeed on a second press.")

    _h2(doc, "Comment the why, never the what")
    _para(doc,
          "The code already says what it does. Comment when there is a decision, a "
          "constraint or a trap: why this and not the obvious alternative, what breaks if "
          "somebody simplifies it, which other file must change with it. src/lib/brand.ts and "
          "src/lib/shop-policy.ts are the reference examples.")

    # ------------------------------------------------------------- the traps
    if data["traps"]:
        _h1(doc, "Traps that cost other people a day")
        _para(doc,
              "Collected from real incidents, and parsed into this document from "
              "Docs/00-start-here.md rather than copied, so it cannot fall behind the table "
              "new joiners actually read.")
        _table(
            doc,
            ["Trap", "What happens", "The truth"],
            [[t[0], t[1], t[2]] for t in data["traps"]],
            widths=[1.9, 2.2, 2.4], size=9,
        )

    # ------------------------------------------------------------- verifying
    _h1(doc, "Testing and deployment")
    _table(
        doc,
        ["Where", "Command", "What it protects"],
        [
            ["Web app", "npx tsc --noEmit ; npm test", f"Types and the parity guards ({c['webTests']} test files)"],
            ["Control plane", "cd platform/api ; npm test", f"Auth, MQTT, households, voice ({c['planeTests']} test files)"],
            ["Mobile", "cd mobile ; npm run typecheck", "tsc plus a dozen static audits"],
            ["Firmware", "python -m platformio run", "It compiles for the board it ships on"],
        ],
        widths=[1.2, 2.2, 3.1],
    )
    _para(doc,
          "The mobile typecheck is not just tsc. It audits navigation targets, device-type "
          "coverage, the command map against firmware behaviour, colour contrast, screen "
          "theming, iOS-only dialogs, swallowed failures and permissions — and each of those "
          "exists because it once shipped a bug. audit-swallowed.js is a ratchet: the number "
          "may fall and must never rise.")
    _para(doc,
          "Deployment differs per deployable and is documented in Docs/09-deployment.md. The "
          "rule that spans all four is the same: confirm the thing you shipped is the thing "
          "that is running. Verify a website by the build sha its health endpoint reports, "
          "never by eye. A control plane that reports commit: unknown was deployed without the "
          "script, and that field exists precisely so a stale container identifies itself "
          "instead of looking like broken hardware.")

    _callout(doc,
             "Verify by running it, not by reading it. Nearly every defect found while this "
             "system was built appeared only when it was executed, and several were in code "
             "written an hour earlier by the person who found them.")

    # ---------------------------------------------------------- the library
    _h1(doc, "The library")
    _para(doc,
          "Do not read these now. Read 00 and 01 today, and the rest when a task calls for "
          "one.")
    _table(
        doc,
        ["Document", "Covers"],
        [[name.replace(".md", ""), title] for name, title in data["docs"]],
        widths=[1.6, 4.9], size=9,
    )

    # ------------------------------------------------------------ first week
    _h1(doc, "Your first week")
    for item in [
        "Work through Docs/00-start-here.md end to end. It ends with four proofs: the site "
        "renders, the tests pass, the API answers, and the firmware compiles. If one fails, "
        "fix it now — it will not get easier once you are also debugging your own code.",
        "Read Docs/01-architecture.md afterwards, not before. It means much more once you "
        "have seen the pieces move.",
        "Pick a task that touches one deployable, and resist anything that needs all four.",
        "Before you open a pull request, ask the parity question: is there another table, "
        "screen or app that also needs to know what I just added? No tool asks it for you.",
        "If you find a document that is wrong, fix it in the same pull request as your code. "
        "These documents are written from the code and only stay true because the person who "
        "notices the drift corrects it.",
    ]:
        _bullet(doc, item)

    doc.save(str(out_path))
    return len(doc.paragraphs)
